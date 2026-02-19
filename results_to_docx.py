#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Convert scanner JSON output into a manager-friendly Word (.docx) report."""

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document


def _severity_key(vuln: dict[str, Any]) -> int:
    order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}
    return order.get(vuln.get("severity", "Info"), 99)


def build_doc(data: dict[str, Any]) -> Document:
    scan_info = data.get("scan_info", {})
    summary = data.get("summary", {})
    vulns = data.get("vulnerabilities", [])

    doc = Document()
    doc.add_heading("Moodle Security Scan Report (Manager Export)", level=0)

    doc.add_paragraph(f"Target: {scan_info.get('target', 'Unknown')}")
    doc.add_paragraph(f"Scan Date: {scan_info.get('timestamp', 'Unknown')}")
    doc.add_paragraph(f"Scanner Version: {scan_info.get('scanner_version', 'Unknown')}")
    doc.add_paragraph(f"Status: {scan_info.get('status', 'Unknown')}")

    doc.add_heading("Scan Summary", level=1)
    doc.add_paragraph(f"Total Vulnerabilities Found: {summary.get('total_vulnerabilities', len(vulns))}")
    doc.add_paragraph(f"Average Risk Score: {summary.get('average_risk_score', 0)}")
    doc.add_paragraph(f"Critical: {summary.get('severity_counts', {}).get('Critical', 0)}")
    doc.add_paragraph(f"High: {summary.get('severity_counts', {}).get('High', 0)}")
    doc.add_paragraph(f"Medium: {summary.get('severity_counts', {}).get('Medium', 0)}")
    doc.add_paragraph(f"Low: {summary.get('severity_counts', {}).get('Low', 0)}")
    doc.add_paragraph(f"Info: {summary.get('severity_counts', {}).get('Info', 0)}")

    doc.add_heading("Remediation Roadmap", level=1)
    doc.add_paragraph("Immediate (0-24h): Fix Critical findings and lock down sensitive endpoints.")
    doc.add_paragraph("Short-term (1-7d): Patch affected Moodle core/plugins and harden config.")
    doc.add_paragraph("Medium-term (1-4w): Add continuous monitoring and recurring scans.")

    doc.add_heading("Safe Validation Note", level=1)
    doc.add_paragraph("This report includes defensive remediation and validation guidance only.")

    doc.add_heading("Detailed Findings", level=1)
    if not vulns:
        doc.add_paragraph("No vulnerabilities found.")
        return doc

    max_evidence_chars = 1000

    for index, vuln in enumerate(sorted(vulns, key=_severity_key), 1):
        title = vuln.get("title", "Unknown vulnerability")
        severity = vuln.get("severity", "Unknown")
        doc.add_heading(f"{index}. [{severity}] {title}", level=2)

        evidence = str(vuln.get("evidence", "N/A"))
        if len(evidence) > max_evidence_chars:
            evidence = evidence[:max_evidence_chars] + "\n...[truncated for readability; see JSON report]"

        doc.add_paragraph(f"Module: {vuln.get('module', 'unknown')}")
        doc.add_paragraph(f"Confidence: {vuln.get('confidence', 'unknown')}")
        doc.add_paragraph(f"Exploitability: {vuln.get('exploitability', 'unknown')}")
        doc.add_paragraph(f"Risk Score: {vuln.get('risk_score', 'unknown')}")
        doc.add_paragraph(f"Description: {vuln.get('description', 'No description provided.')}")
        doc.add_paragraph(f"URL: {vuln.get('url', 'N/A')}")
        doc.add_paragraph(f"Evidence: {evidence}")

        if vuln.get("cve"):
            cve = vuln.get("cve")
            doc.add_paragraph(f"CVE: {cve}")
            doc.add_paragraph(f"Reference: https://nvd.nist.gov/vuln/detail/{cve}")

        if vuln.get("cwe"):
            cwe = str(vuln.get("cwe"))
            doc.add_paragraph(f"CWE: {cwe}")
            doc.add_paragraph(f"Reference: https://cwe.mitre.org/data/definitions/{cwe.replace('CWE-', '')}.html")

        if vuln.get("payload"):
            doc.add_paragraph("Payload Details: omitted in manager report")

        doc.add_paragraph(f"Remediation: {vuln.get('remediation', 'No remediation provided.')}")

        doc.add_paragraph("Fix Steps (Actionable):")
        for step in [
            "Assign an owner and change window.",
            "Patch/update impacted Moodle component or plugin.",
            "Disable/remove vulnerable component if no patch is available.",
            "Restrict access/authentication for sensitive endpoints.",
            "Re-test and document closure evidence.",
        ]:
            doc.add_paragraph(step, style="List Number")

        doc.add_paragraph("Safe Validation Checklist:")
        for step in [
            "Confirm endpoint no longer behaves as vulnerable.",
            "Confirm production errors do not expose sensitive internals.",
            "Re-run scanner module and verify finding is gone/downgraded.",
        ]:
            doc.add_paragraph(step, style="List Bullet")

    return doc


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert scanner JSON output into manager DOCX report")
    parser.add_argument("-i", "--input", required=True, help="Input JSON file (e.g., results.json)")
    parser.add_argument("-o", "--output", default="manager_report.docx", help="Output DOCX file")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    with input_path.open("r", encoding="utf-8") as source:
        data = json.load(source)

    doc = build_doc(data)
    doc.save(str(output_path))
    print(f"Report written to {output_path}")


if __name__ == "__main__":
    main()
